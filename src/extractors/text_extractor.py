"""Text extraction service for various document formats."""

import logging
import io
import os
import tempfile
from typing import Optional, Dict, Any
from abc import ABC, abstractmethod
import chardet
import pandas as pd
from PIL import Image
import pytesseract
import PyPDF2
import pdfplumber
from docx import Document
from openpyxl import load_workbook
from config import settings

logger = logging.getLogger(__name__)


class BaseExtractor(ABC):
    """Base class for text extractors."""
    
    @abstractmethod
    def extract(self, content: bytes, file_name: str) -> str:
        """Extract text from file content."""
        pass
    
    @abstractmethod
    def can_handle(self, file_extension: str) -> bool:
        """Check if this extractor can handle the given file extension."""
        pass


class TxtExtractor(BaseExtractor):
    """Extracts text from plain text files."""
    
    def can_handle(self, file_extension: str) -> bool:
        return file_extension in ['.txt', '.log', '.md', '.readme']
    
    def extract(self, content: bytes, file_name: str) -> str:
        """Extract text from plain text files with encoding detection."""
        try:
            # Detect encoding
            encoding_result = chardet.detect(content)
            encoding = encoding_result.get('encoding', 'utf-8')
            
            if encoding is None:
                encoding = 'utf-8'
            
            # Try to decode with detected encoding
            try:
                text = content.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                # Fallback to utf-8 with error handling
                text = content.decode('utf-8', errors='replace')
            
            logger.debug(f"Extracted {len(text)} characters from {file_name}")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_name}: {e}")
            return ""


class CsvExtractor(BaseExtractor):
    """Extracts text content from CSV files."""
    
    def can_handle(self, file_extension: str) -> bool:
        return file_extension in ['.csv', '.tsv']
    
    def extract(self, content: bytes, file_name: str) -> str:
        """Extract text from CSV files."""
        try:
            # Detect encoding
            encoding_result = chardet.detect(content)
            encoding = encoding_result.get('encoding', 'utf-8')
            
            if encoding is None:
                encoding = 'utf-8'
            
            # Read CSV into DataFrame
            try:
                df = pd.read_csv(
                    io.BytesIO(content),
                    encoding=encoding,
                    on_bad_lines='skip'
                )
            except UnicodeDecodeError:
                df = pd.read_csv(
                    io.BytesIO(content),
                    encoding='utf-8',
                    on_bad_lines='skip',
                    errors='replace'
                )
            
            # Convert all data to string and join
            text_parts = []
            
            # Add column headers
            text_parts.extend(df.columns.tolist())
            
            # Add all cell values
            for _, row in df.iterrows():
                for value in row.values:
                    if pd.notna(value):
                        text_parts.append(str(value))
            
            text = ' '.join(text_parts)
            logger.debug(f"Extracted {len(text)} characters from CSV {file_name}")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from CSV {file_name}: {e}")
            return ""


class PdfExtractor(BaseExtractor):
    """Extracts text from PDF files using multiple methods."""
    
    def can_handle(self, file_extension: str) -> bool:
        return file_extension == '.pdf'
    
    def extract(self, content: bytes, file_name: str) -> str:
        """Extract text from PDF files using pdfplumber first, then PyPDF2 as fallback."""
        text = self._extract_with_pdfplumber(content, file_name)
        
        if not text.strip():
            logger.debug(f"pdfplumber failed for {file_name}, trying PyPDF2")
            text = self._extract_with_pypdf2(content, file_name)
        
        logger.debug(f"Extracted {len(text)} characters from PDF {file_name}")
        return text
    
    def _extract_with_pdfplumber(self, content: bytes, file_name: str) -> str:
        """Extract text using pdfplumber (better for complex layouts)."""
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                
                return '\n'.join(text_parts)
                
        except Exception as e:
            logger.debug(f"pdfplumber extraction failed for {file_name}: {e}")
            return ""
    
    def _extract_with_pypdf2(self, content: bytes, file_name: str) -> str:
        """Extract text using PyPDF2 as fallback."""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
            text_parts = []
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            logger.debug(f"PyPDF2 extraction failed for {file_name}: {e}")
            return ""


class ImageExtractor(BaseExtractor):
    """Extracts text from images using OCR (Tesseract)."""
    
    def can_handle(self, file_extension: str) -> bool:
        return file_extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']
    
    def extract(self, content: bytes, file_name: str) -> str:
        """Extract text from images using Tesseract OCR."""
        try:
            # Configure Tesseract path if specified
            if settings.tesseract_path and os.path.exists(settings.tesseract_path):
                pytesseract.pytesseract.tesseract_cmd = settings.tesseract_path
            
            # Open image from bytes
            image = Image.open(io.BytesIO(content))
            
            # Perform OCR
            text = pytesseract.image_to_string(image, lang='eng')
            
            logger.debug(f"Extracted {len(text)} characters from image {file_name}")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from image {file_name}: {e}")
            return ""


class DocxExtractor(BaseExtractor):
    """Extracts text from Word documents."""
    
    def can_handle(self, file_extension: str) -> bool:
        return file_extension in ['.docx', '.doc']
    
    def extract(self, content: bytes, file_name: str) -> str:
        """Extract text from Word documents."""
        try:
            doc = Document(io.BytesIO(content))
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            text = '\n'.join(text_parts)
            logger.debug(f"Extracted {len(text)} characters from Word document {file_name}")
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from Word document {file_name}: {e}")
            return ""


class ExcelExtractor(BaseExtractor):
    """Extracts text from Excel files."""
    
    def can_handle(self, file_extension: str) -> bool:
        return file_extension in ['.xlsx', '.xls']
    
    def extract(self, content: bytes, file_name: str) -> str:
        """Extract text from Excel files."""
        try:
            # Save content to temporary file (openpyxl needs file path)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as temp_file:
                temp_file.write(content)
                temp_file_path = temp_file.name
            
            try:
                workbook = load_workbook(temp_file_path, data_only=True)
                text_parts = []
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    
                    # Add sheet name
                    text_parts.append(f"Sheet: {sheet_name}")
                    
                    # Extract all cell values
                    for row in sheet.iter_rows(values_only=True):
                        for cell_value in row:
                            if cell_value is not None:
                                text_parts.append(str(cell_value))
                
                text = ' '.join(text_parts)
                logger.debug(f"Extracted {len(text)} characters from Excel {file_name}")
                return text
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
        except Exception as e:
            logger.error(f"Error extracting text from Excel {file_name}: {e}")
            return ""


class TextExtractorService:
    """Main service for extracting text from various document types."""
    
    def __init__(self):
        """Initialize the text extractor service with all available extractors."""
        self.extractors = [
            TxtExtractor(),
            CsvExtractor(),
            PdfExtractor(),
            ImageExtractor(),
            DocxExtractor(),
            ExcelExtractor()
        ]
        logger.info(f"Initialized TextExtractorService with {len(self.extractors)} extractors")
    
    def extract_text(self, content: bytes, file_name: str) -> str:
        """Extract text from file content based on file extension."""
        if not content:
            logger.warning(f"Empty content for file: {file_name}")
            return ""
        
        # Get file extension
        file_extension = os.path.splitext(file_name)[1].lower()
        
        if not file_extension:
            logger.warning(f"No file extension found for: {file_name}")
            return ""
        
        # Find appropriate extractor
        for extractor in self.extractors:
            if extractor.can_handle(file_extension):
                try:
                    text = extractor.extract(content, file_name)
                    if text:
                        logger.info(f"Successfully extracted text from {file_name} using {extractor.__class__.__name__}")
                        return text
                    else:
                        logger.warning(f"No text extracted from {file_name} using {extractor.__class__.__name__}")
                        return ""
                except Exception as e:
                    logger.error(f"Error in {extractor.__class__.__name__} for {file_name}: {e}")
                    return ""
        
        logger.warning(f"No extractor found for file extension: {file_extension}")
        return ""
    
    def get_supported_extensions(self) -> list:
        """Get list of all supported file extensions."""
        extensions = set()
        for extractor in self.extractors:
            # This is a simplified check - in practice, you'd need to implement
            # a get_supported_extensions method in each extractor
            pass
        return settings.supported_extensions_list
    
    def can_process_file(self, file_name: str) -> bool:
        """Check if the file can be processed by any extractor."""
        file_extension = os.path.splitext(file_name)[1].lower()
        return any(extractor.can_handle(file_extension) for extractor in self.extractors) 